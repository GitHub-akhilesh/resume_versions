import os
import re
from lxml import html
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from playwright.sync_api import sync_playwright

# Define paths
HTML_FILE = "resume_preview.html"
DOWNLOADS_DIR = "downloads"

# Version mapping: key in HTML id -> file name prefix
VERSIONS = {
    "version1-enhancv": "Akhilesh_Mishra_MERN_Stack_Enhancv",
    "version2-enhancv": "Akhilesh_Mishra_Java_Full_Stack_Enhancv",
    "version3-enhancv": "Akhilesh_Mishra_Python_Full_Stack_Enhancv",
    "version4-enhancv": "Akhilesh_Mishra_Software_Engineer_Enhancv",
    "version1-modern": "Akhilesh_Mishra_MERN_Stack",
    "version2-modern": "Akhilesh_Mishra_Java_Full_Stack",
    "version3-modern": "Akhilesh_Mishra_Python_Full_Stack",
    "version4-modern": "Akhilesh_Mishra_Software_Engineer",
    
    "version1-classic": "Akhilesh_Mishra_MERN_Stack_Classic",
    "version2-classic": "Akhilesh_Mishra_Java_Full_Stack_Classic",
    "version3-classic": "Akhilesh_Mishra_Python_Full_Stack_Classic",
    "version4-classic": "Akhilesh_Mishra_Software_Engineer_Classic",
    
    "version1-minimal": "Akhilesh_Mishra_MERN_Stack_Minimal",
    "version2-minimal": "Akhilesh_Mishra_Java_Full_Stack_Minimal",
    "version3-minimal": "Akhilesh_Mishra_Python_Full_Stack_Minimal",
    "version4-minimal": "Akhilesh_Mishra_Software_Engineer_Minimal",
    
    "version1-executive": "Akhilesh_Mishra_MERN_Stack_Executive",
    "version2-executive": "Akhilesh_Mishra_Java_Full_Stack_Executive",
    "version3-executive": "Akhilesh_Mishra_Python_Full_Stack_Executive",
    "version4-executive": "Akhilesh_Mishra_Software_Engineer_Executive",
}

# Version source markdown file mapping
MD_FILES = {
    "version1": "version1_mern.md",
    "version2": "version2_java.md",
    "version3": "version3_python.md",
    "version4": "version4_combined.md",
}

def setup_directories():
    if not os.path.exists(DOWNLOADS_DIR):
        os.makedirs(DOWNLOADS_DIR)
        print(f"Created directory: {DOWNLOADS_DIR}")

def extract_latex():
    """Extracts LaTeX source from markdown files and writes to .tex files."""
    for ver, md_filename in MD_FILES.items():
        if not os.path.exists(md_filename):
            print(f"Warning: {md_filename} not found!")
            continue
            
        with open(md_filename, "r", encoding="utf-8") as f:
            content = f.read()
            
        # Search for ```latex ... ``` block
        match = re.search(r"```latex\\s*(.*?)\\s*```", content, re.DOTALL)
        if not match:
            match = re.search(r"```latex\s*(.*?)\s*```", content, re.DOTALL)
            
        if match:
            latex_source = match.group(1).strip()
            
            # 1. Modern tex
            modern_path = os.path.join(DOWNLOADS_DIR, f"{VERSIONS[ver + '-modern']}.tex")
            with open(modern_path, "w", encoding="utf-8") as out:
                out.write(latex_source)
            print(f"Generated LaTeX (Modern): {modern_path}")
            
            # 2. Classic tex (royal blue)
            classic_source = latex_source.replace("{0c4f6b}", "{0056b3}")
            classic_path = os.path.join(DOWNLOADS_DIR, f"{VERSIONS[ver + '-classic']}.tex")
            with open(classic_path, "w", encoding="utf-8") as out:
                out.write(classic_source)
            print(f"Generated LaTeX (Classic): {classic_path}")

            # 3. Minimal tex (charcoal / dark gray)
            minimal_source = latex_source.replace("{0c4f6b}", "{1a1a1a}")
            # replace helvet with times / georgia style serif fonts if needed, or keep clean
            minimal_path = os.path.join(DOWNLOADS_DIR, f"{VERSIONS[ver + '-minimal']}.tex")
            with open(minimal_path, "w", encoding="utf-8") as out:
                out.write(minimal_source)
            print(f"Generated LaTeX (Minimal): {minimal_path}")

            # 4. Executive tex (slate gray / navy)
            exec_source = latex_source.replace("{0c4f6b}", "{1e293b}")
            exec_path = os.path.join(DOWNLOADS_DIR, f"{VERSIONS[ver + '-executive']}.tex")
            with open(exec_path, "w", encoding="utf-8") as out:
                out.write(exec_source)
            print(f"Generated LaTeX (Executive): {exec_path}")

            # 5. Enhancv tex (blue accent)
            enhancv_source = latex_source.replace("{0c4f6b}", "{007bb6}")
            enhancv_path = os.path.join(DOWNLOADS_DIR, f"{VERSIONS[ver + '-enhancv']}.tex")
            with open(enhancv_path, "w", encoding="utf-8") as out:
                out.write(enhancv_source)
            print(f"Generated LaTeX (Enhancv): {enhancv_path}")

def add_p_border_bottom(paragraph, color_hex="0C4F6B", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     rf'<w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color_hex}"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)

def generate_word_docs():
    """Parses HTML and generates beautifully formatted Word (.docx) documents."""
    if not os.path.exists(HTML_FILE):
        print(f"Error: {HTML_FILE} not found!")
        return

    with open(HTML_FILE, "r", encoding="utf-8") as f:
        html_content = f.read()

    tree = html.fromstring(html_content)

    for ver_id, filename_prefix in VERSIONS.items():
        doc = Document()
        
        # Setup margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

        sheet = tree.xpath(f'//div[@id="{ver_id}"]')
        if not sheet:
            print(f"Warning: Container {ver_id} not found in HTML!")
            continue
        
        sheet = sheet[0]
        
        # Theme configuration
        is_modern = ver_id.endswith("-modern")
        is_classic = ver_id.endswith("-classic")
        is_minimal = ver_id.endswith("-minimal")
        is_executive = ver_id.endswith("-executive")
        is_enhancv = ver_id.endswith("-enhancv")
        
        font_name = "Georgia" if is_minimal else "Arial"
        
        if is_modern:
            primary_color_rgb = RGBColor(0x0c, 0x4f, 0x6b)
            border_hex = "0C4F6B"
        elif is_classic:
            primary_color_rgb = RGBColor(0x00, 0x56, 0xb3)
            border_hex = "0056B3"
        elif is_minimal:
            primary_color_rgb = RGBColor(0x1a, 0x1a, 0x1a)
            border_hex = "1A1A1A"
        elif is_executive:
            primary_color_rgb = RGBColor(0x1e, 0x29, 0x3b)
            border_hex = "1E293B"
        else: # enhancv
            primary_color_rgb = RGBColor(0x00, 0x7b, 0xb6)
            border_hex = "007BB6"

        header_align = WD_ALIGN_PARAGRAPH.LEFT if (is_modern or is_executive or is_enhancv) else WD_ALIGN_PARAGRAPH.CENTER

        if is_enhancv:
            # Render Enhancv full-width header at the top
            header = sheet.xpath('.//div[@class="header"]')[0]
            name = header.xpath('.//div[@class="name"]/text()')[0]
            title = header.xpath('.//div[@class="title"]/text()')[0]
            
            p_name = doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(0)
            p_name.paragraph_format.space_after = Pt(2)
            run = p_name.add_run(name.strip())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
            
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(4)
            run = p_title.add_run(title.strip())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(11)
            run.font.color.rgb = primary_color_rgb
            
            # Contact Info
            contact_lines = header.xpath('.//div[@class="contact-info"]/span//text() | .//div[@class="contact-info"]/span/a//text()')
            contact_lines = [l.strip() for l in contact_lines if l.strip() and l.strip() != "|"]
            p_contact = doc.add_paragraph()
            p_contact.paragraph_format.space_before = Pt(0)
            p_contact.paragraph_format.space_after = Pt(8)
            p_contact.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_p_border_bottom(p_contact, color_hex="1A1A1A", size="8")
            
            first = True
            for line in contact_lines:
                if not first:
                    p_contact.add_run("   |   ")
                first = False
                run = p_contact.add_run(line)
                run.font.name = font_name
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

        if is_executive or is_enhancv:
            # 2-COLUMN LAYOUT
            # 2-COLUMN LAYOUT
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            
            if is_executive:
                table.rows[0].cells[0].width = Inches(2.3)
                table.rows[0].cells[1].width = Inches(5.2)
                sidebar_cell = table.cell(0, 0)
                main_cell = table.cell(0, 1)
            else: # enhancv
                table.rows[0].cells[0].width = Inches(5.2)
                table.rows[0].cells[1].width = Inches(2.3)
                main_cell = table.cell(0, 0)
                sidebar_cell = table.cell(0, 1)
            
            # Sidebar background shading for Executive only
            tcPr_side = sidebar_cell._tc.get_or_add_tcPr()
            if is_executive:
                shd_side = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F8FAFC"/>')
                tcPr_side.append(shd_side)
            
            # Cell Margins for sidebar
            tcMar_side = parse_xml(r'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                   r'<w:top w:w="240" w:type="dxa"/>'
                                   r'<w:bottom w:w="240" w:type="dxa"/>'
                                   r'<w:left w:w="180" w:type="dxa"/>'
                                   r'<w:right w:w="180" w:type="dxa"/>'
                                   r'</w:tcMar>')
            tcPr_side.append(tcMar_side)

            # Cell Margins for main cell
            tcPr_main = main_cell._tc.get_or_add_tcPr()
            tcMar_main = parse_xml(r'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                   r'<w:top w:w="240" w:type="dxa"/>'
                                   r'<w:bottom w:w="240" w:type="dxa"/>'
                                   r'<w:left w:w="240" w:type="dxa"/>'
                                   r'<w:right w:w="240" w:type="dxa"/>'
                                   r'</w:tcMar>')
            tcPr_main.append(tcMar_main)

            if is_executive:
                # Sidebar Content Building (Original Executive style)
                sidebar_el = sheet.xpath('.//div[@class="executive-sidebar"]')[0]
                
                # Name Block
                name = sidebar_el.xpath('.//div[@class="name"]/text()')
                title = sidebar_el.xpath('.//div[@class="title"]/text()')
                if name:
                    p_name = sidebar_cell.paragraphs[0]
                    p_name.paragraph_format.space_before = Pt(0)
                    p_name.paragraph_format.space_after = Pt(2)
                    run = p_name.add_run(" ".join(name).strip())
                    run.bold = True
                    run.font.name = font_name
                    run.font.size = Pt(18)
                    run.font.color.rgb = primary_color_rgb
                
                if title:
                    p_title = sidebar_cell.add_paragraph()
                    p_title.paragraph_format.space_before = Pt(0)
                    p_title.paragraph_format.space_after = Pt(8)
                    run = p_title.add_run(title[0].strip())
                    run.bold = True
                    run.font.name = font_name
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

                # Contact
                sidebar_cell.add_paragraph().add_run("CONTACT").bold = True
                add_p_border_bottom(sidebar_cell.paragraphs[-1], color_hex=border_hex, size="4")
                sidebar_cell.paragraphs[-1].paragraph_format.space_before = Pt(8)
                sidebar_cell.paragraphs[-1].paragraph_format.space_after = Pt(4)
                
                contact_lines = sidebar_el.xpath('.//div[@class="contact-info"]/span//text() | .//div[@class="contact-info"]/span/a//text()')
                contact_lines = [l.strip() for l in contact_lines if l.strip() and l.strip() != "|"]
                for line in contact_lines:
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(line)
                    run.font.name = font_name
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

                # Skills
                p_sk = sidebar_cell.add_paragraph()
                p_sk.paragraph_format.space_before = Pt(12)
                p_sk.paragraph_format.space_after = Pt(4)
                p_sk.add_run("CORE SKILLS").bold = True
                add_p_border_bottom(p_sk, color_hex=border_hex, size="4")
                
                skills_items = sidebar_el.xpath('.//div[@class="sidebar-skills"]/div[@class="edu-item"]')
                for sk in skills_items:
                    strong_tag = sk.find('strong')
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if strong_tag is not None:
                        run_strong = p.add_run(strong_tag.text_content().strip() + "\n")
                        run_strong.bold = True
                        run_strong.font.name = font_name
                        run_strong.font.size = Pt(8.5)
                        run_strong.font.color.rgb = primary_color_rgb
                        
                        rest_text = sk.text_content()[len(strong_tag.text_content()):].strip()
                        run_rest = p.add_run(rest_text)
                        run_rest.font.name = font_name
                        run_rest.font.size = Pt(8)
                        run_rest.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

                # Education
                p_edu = sidebar_cell.add_paragraph()
                p_edu.paragraph_format.space_before = Pt(12)
                p_edu.paragraph_format.space_after = Pt(4)
                p_edu.add_run("EDUCATION").bold = True
                add_p_border_bottom(p_edu, color_hex=border_hex, size="4")
                
                edu_el = sidebar_el.xpath('.//div[contains(@class,"edu-item") and not(parent::div[@class="sidebar-skills"])]')
                if edu_el:
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(edu_el[0].text_content().strip())
                    run.font.name = font_name
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

                # Certifications
                p_cert = sidebar_cell.add_paragraph()
                p_cert.paragraph_format.space_before = Pt(12)
                p_cert.paragraph_format.space_after = Pt(4)
                p_cert.add_run("CERTIFICATIONS").bold = True
                add_p_border_bottom(p_cert, color_hex=border_hex, size="4")
                
                cert_el = sidebar_el.xpath('.//div[@class="cert-list"]')
                if cert_el:
                    cert_lines = [l.strip() for l in cert_el[0].text_content().split('\n') if l.strip()]
                    for line in cert_lines:
                        p = sidebar_cell.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        run = p.add_run(line)
                        run.font.name = font_name
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            else:
                # Enhancv Sidebar (Right column) Content Building
                sidebar_el = sheet.xpath('.//div[@class="enhancv-right"]')[0]
                
                # Profile Summary
                p_sum = sidebar_cell.paragraphs[0]
                p_sum.paragraph_format.space_before = Pt(0)
                p_sum.paragraph_format.space_after = Pt(4)
                p_sum.add_run("SUMMARY").bold = True
                add_p_border_bottom(p_sum, color_hex=border_hex, size="4")
                
                sum_text = sidebar_el.xpath('.//div[@class="summary-box"]/div[@class="summary-text"]/text()')
                if sum_text:
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(8)
                    run = p.add_run(sum_text[0].strip())
                    run.font.name = font_name
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

                # Education
                p_edu = sidebar_cell.add_paragraph()
                p_edu.paragraph_format.space_before = Pt(8)
                p_edu.paragraph_format.space_after = Pt(4)
                p_edu.add_run("EDUCATION").bold = True
                add_p_border_bottom(p_edu, color_hex=border_hex, size="4")
                
                edu_box = sidebar_el.xpath('.//div[@class="education-bar"]')[0]
                p = sidebar_cell.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(edu_box.text_content().strip())
                run.font.name = font_name
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

                # Skills
                p_sk = sidebar_cell.add_paragraph()
                p_sk.paragraph_format.space_before = Pt(12)
                p_sk.paragraph_format.space_after = Pt(4)
                p_sk.add_run("SKILLS").bold = True
                add_p_border_bottom(p_sk, color_hex=border_hex, size="4")
                
                sk_sections = sidebar_el.xpath('.//div[@class="skills-section"]')
                for section in sk_sections:
                    sec_title = section.xpath('.//div[@class="skills-sec-title"]/text()')[0].strip()
                    tags = section.xpath('.//span[@class="skill-tag"]/text()')
                    tags_str = ", ".join([t.strip() for t in tags])
                    
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    run_title = p.add_run(sec_title + ": ")
                    run_title.bold = True
                    run_title.font.name = font_name
                    run_title.font.size = Pt(8.5)
                    run_title.font.color.rgb = primary_color_rgb
                    
                    run_tags = p.add_run(tags_str)
                    run_tags.font.name = font_name
                    run_tags.font.size = Pt(8)
                    run_tags.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

                # Certifications
                p_cert = sidebar_cell.add_paragraph()
                p_cert.paragraph_format.space_before = Pt(12)
                p_cert.paragraph_format.space_after = Pt(4)
                p_cert.add_run("CERTIFICATIONS").bold = True
                add_p_border_bottom(p_cert, color_hex=border_hex, size="4")
                
                cert_box = sidebar_el.xpath('.//div[@class="cert-list"]')[0]
                cert_lines = [l.strip() for l in cert_box.text_content().split('\n') if l.strip()]
                for line in cert_lines:
                    p = sidebar_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(line)
                    run.font.name = font_name
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)

                # Achievements
                p_ach = sidebar_cell.add_paragraph()
                p_ach.paragraph_format.space_before = Pt(12)
                p_ach.paragraph_format.space_after = Pt(4)
                p_ach.add_run("ACHIEVEMENTS").bold = True
                add_p_border_bottom(p_ach, color_hex=border_hex, size="4")
                
                ach_items = sidebar_el.xpath('.//ul[@class="achievements-list"]/li/text()')
                for ach in ach_items:
                    p = sidebar_cell.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(ach.strip())
                    run.font.name = font_name
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)


            # Main Content Cell Building
            if is_executive:
                main_el = sheet.xpath('.//div[@class="executive-main"]')[0]
            else: # enhancv
                main_el = sheet.xpath('.//div[@class="enhancv-left"]')[0]
                
            main_children = main_el.xpath('./*')
            
            first_p = True
            for child in main_children:
                c_class = child.get('class', '')
                
                if c_class == 'section-title':
                    p = main_cell.paragraphs[0] if first_p else main_cell.add_paragraph()
                    first_p = False
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(child.text_content().strip().upper())
                    run.bold = True
                    run.font.name = font_name
                    run.font.size = Pt(11)
                    run.font.color.rgb = primary_color_rgb
                    add_p_border_bottom(p, color_hex=border_hex, size="6")
                    
                elif c_class == 'summary-box':
                    p = main_cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(child.text_content().strip())
                    run.font.name = font_name
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                    
                elif c_class == 'experience-container':
                    items = child.xpath('.//div[@class="experience-item"]')
                    for exp in items:
                        role = exp.xpath('.//span[@class="role-title"]/text()')
                        company = exp.xpath('.//span[@class="company-name"]/text()')
                        desc = exp.xpath('.//div[@class="company-desc"]/text()')
                        meta = exp.xpath('.//div[@class="meta-line"]/text()')
                        bullets = exp.xpath('.//ul/li | .//ul[@class="bullet-list"]/li')
                        
                        p_role = main_cell.add_paragraph()
                        p_role.paragraph_format.space_before = Pt(4)
                        p_role.paragraph_format.space_after = Pt(1)
                        if role:
                            run = p_role.add_run(role[0].strip())
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(10.5)
                            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
                        if role and company:
                            p_role.add_run(" - ")
                        if company:
                            run = p_role.add_run(company[0].strip())
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                            
                        if desc:
                            p_desc = main_cell.add_paragraph()
                            p_desc.paragraph_format.space_before = Pt(0)
                            p_desc.paragraph_format.space_after = Pt(1)
                            run = p_desc.add_run(desc[0].strip())
                            run.font.name = font_name
                            run.font.size = Pt(8.5)
                            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                            run.italic = True
                            
                        if meta:
                            p_meta = main_cell.add_paragraph()
                            p_meta.paragraph_format.space_before = Pt(0)
                            p_meta.paragraph_format.space_after = Pt(3)
                            run = p_meta.add_run(meta[0].strip())
                            run.font.name = font_name
                            run.font.size = Pt(8.5)
                            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                            
                        for li in bullets:
                            p_bullet = main_cell.add_paragraph(style='List Bullet')
                            p_bullet.paragraph_format.space_before = Pt(0)
                            p_bullet.paragraph_format.space_after = Pt(1)
                            p_bullet.paragraph_format.left_indent = Inches(0.15)
                            
                            # Parse strong text inside bullets
                            for node in li.xpath('./* | text()'):
                                if isinstance(node, str):
                                    run = p_bullet.add_run(node)
                                    run.font.name = font_name
                                    run.font.size = Pt(9.5)
                                    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                                else:
                                    run = p_bullet.add_run(node.text_content())
                                    run.font.name = font_name
                                    run.font.size = Pt(9.5)
                                    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                                    if node.tag == 'strong':
                                        run.bold = True

                elif c_class == 'projects-grid':
                    projs = child.xpath('.//div[@class="project-card"]')
                    for proj in projs:
                        title = proj.xpath('.//div[@class="project-title"]/text()')
                        desc = proj.xpath('.//div[@class="project-desc"]/text()')
                        stack = proj.xpath('.//div[@class="project-stack"]//text()')
                        
                        p_title = main_cell.add_paragraph()
                        p_title.paragraph_format.space_before = Pt(4)
                        p_title.paragraph_format.space_after = Pt(1)
                        if title:
                            run = p_title.add_run(title[0].strip())
                            run.bold = True
                            run.font.name = font_name
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
                            
                        if desc:
                            p_desc = main_cell.add_paragraph()
                            p_desc.paragraph_format.space_before = Pt(0)
                            p_desc.paragraph_format.space_after = Pt(1)
                            run = p_desc.add_run(desc[0].strip())
                            run.font.name = font_name
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                            
                        if stack:
                            p_stack = main_cell.add_paragraph()
                            p_stack.paragraph_format.space_before = Pt(0)
                            p_stack.paragraph_format.space_after = Pt(4)
                            run = p_stack.add_run("".join(stack).strip())
                            run.font.name = font_name
                            run.font.size = Pt(8.5)
                            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
                            run.italic = True
                            
                elif c_class == 'achievements-list' or child.tag == 'ul':
                    for li in child.xpath('.//li'):
                        p = main_cell.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.left_indent = Inches(0.15)
                        run = p.add_run(li.text_content().strip())
                        run.font.name = font_name
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                        
        else:
            # STANDARD 1-COLUMN DOCUMENT LAYOUT (MODERN, CLASSIC, MINIMAL)
            # 1. Header Information
            header = sheet.xpath('.//div[@class="header"]')
            if header:
                header = header[0]
                name = header.xpath('.//div[@class="name"]/text()')
                title = header.xpath('.//div[@class="title"]/text()')
                
                if name:
                    p_name = doc.add_paragraph()
                    p_name.alignment = header_align
                    p_name.paragraph_format.space_before = Pt(0)
                    p_name.paragraph_format.space_after = Pt(2)
                    run = p_name.add_run(name[0].strip())
                    run.font.name = font_name
                    run.font.size = Pt(22 if is_modern else 20)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x0c, 0x4f, 0x6b) if is_modern else RGBColor(0x1a, 0x1a, 0x1a)

                if title:
                    p_title = doc.add_paragraph()
                    p_title.alignment = header_align
                    p_title.paragraph_format.space_before = Pt(0)
                    p_title.paragraph_format.space_after = Pt(4)
                    run = p_title.add_run(title[0].strip())
                    run.font.name = font_name
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.color.rgb = primary_color_rgb

                contact_div = header.xpath('.//div[@class="contact-info"]')
                if contact_div:
                    contact_div = contact_div[0]
                    parts = []
                    for node in contact_div.xpath('.//span/text() | .//a/text()'):
                        text = node.strip()
                        if text and text != "|":
                            parts.append(text)
                    
                    contact_str = "   |   ".join(parts)
                    p_contact = doc.add_paragraph()
                    p_contact.alignment = header_align
                    p_contact.paragraph_format.space_before = Pt(0)
                    p_contact.paragraph_format.space_after = Pt(10)
                    run = p_contact.add_run(contact_str)
                    run.font.name = font_name
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # 2. Body Elements
            body_elements = sheet.xpath('./*')
            skip_tags = ['ats-badge', 'header']
            
            i = 0
            while i < len(body_elements):
                el = body_elements[i]
                el_class = el.get('class', '')

                if any(s in el_class for s in skip_tags) or el_class == 'section-divider':
                    i += 1
                    continue

                if el_class == 'section-title':
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.keep_with_next = True
                    run = p.add_run(el.text_content().strip().upper())
                    run.font.name = font_name
                    run.font.size = Pt(11.5)
                    run.bold = True
                    run.font.color.rgb = primary_color_rgb
                    add_p_border_bottom(p, color_hex=border_hex, size="8")

                elif el_class == 'summary-box':
                    summary_text = el.xpath('.//div[@class="summary-text"]/text()')
                    if not summary_text:
                        summary_text = [el.text_content()]
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(summary_text[0].strip())
                    run.font.name = font_name
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

                elif el_class == 'experience-container':
                    items = el.xpath('.//div[@class="experience-item"]')
                    for item in items:
                        role = item.xpath('.//span[@class="role-title"]/text()')
                        company = item.xpath('.//span[@class="company-name"]/text()')
                        desc = item.xpath('.//div[@class="company-desc"]/text()')
                        left_text_nodes = item.xpath('.//div[@class="experience-left"]//text()')
                        bullets = item.xpath('.//ul[@class="bullet-list"]/li | .//ul/li')

                        left_parts = [node.strip() for node in left_text_nodes if node.strip()]
                        meta_text = "   |   ".join(left_parts)
                        if desc:
                            meta_text += f"   |   {desc[0].strip()}"

                        if role or company:
                            p_role = doc.add_paragraph()
                            p_role.paragraph_format.space_before = Pt(6)
                            p_role.paragraph_format.space_after = Pt(1)
                            p_role.paragraph_format.keep_with_next = True
                            
                            if role:
                                run_role = p_role.add_run(role[0].strip())
                                run_role.bold = True
                                run_role.font.name = font_name
                                run_role.font.size = Pt(10.5)
                                run_role.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                            
                            if role and company:
                                p_role.add_run("   |   ")
                                
                            if company:
                                run_comp = p_role.add_run(company[0].strip())
                                run_comp.bold = True
                                run_comp.font.name = font_name
                                run_comp.font.size = Pt(10.5)
                                run_comp.font.color.rgb = RGBColor(0x00, 0x56, 0xb3)

                        if meta_text:
                            p_meta = doc.add_paragraph()
                            p_meta.paragraph_format.space_before = Pt(0)
                            p_meta.paragraph_format.space_after = Pt(3)
                            p_meta.paragraph_format.keep_with_next = True
                            run_meta = p_meta.add_run(meta_text)
                            run_meta.font.name = font_name
                            run_meta.font.size = Pt(9.5)
                            run_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                            run_meta.italic = True

                        for li in bullets:
                            p_bullet = doc.add_paragraph(style='List Bullet')
                            p_bullet.paragraph_format.space_before = Pt(0)
                            p_bullet.paragraph_format.space_after = Pt(2)
                            p_bullet.paragraph_format.left_indent = Inches(0.2)
                            
                            li_children = li.xpath('./* | text()')
                            for node in li_children:
                                if isinstance(node, str):
                                    text = node
                                    if text:
                                        run = p_bullet.add_run(text)
                                        run.font.name = font_name
                                        run.font.size = Pt(10)
                                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                else:
                                    text = node.text_content()
                                    if text:
                                        run = p_bullet.add_run(text)
                                        run.font.name = font_name
                                        run.font.size = Pt(10)
                                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                        if node.tag == 'strong':
                                            run.bold = True
                                            run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

                elif el_class == 'experience-item':
                    role = el.xpath('.//div[@class="role-title"]/text()')
                    company = el.xpath('.//div[@class="company-name"]/text()')
                    company_desc = el.xpath('.//div[@class="company-desc"]/text()')
                    meta = el.xpath('.//div[@class="meta-line"]/text()')
                    bullets = el.xpath('.//ul[@class="bullet-list"]/li | .//ul/li')

                    if role or company:
                        p_role = doc.add_paragraph()
                        p_role.paragraph_format.space_before = Pt(6)
                        p_role.paragraph_format.space_after = Pt(1)
                        p_role.paragraph_format.keep_with_next = True
                        
                        if role:
                            run_role = p_role.add_run(role[0].strip())
                            run_role.bold = True
                            run_role.font.name = font_name
                            run_role.font.size = Pt(10.5)
                            run_role.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                        
                        if role and company:
                            p_role.add_run("   |   ")
                            
                        if company:
                            run_comp = p_role.add_run(company[0].strip())
                            run_comp.bold = True
                            run_comp.font.name = font_name
                            run_comp.font.size = Pt(10.5)
                            run_comp.font.color.rgb = primary_color_rgb

                    if company_desc:
                        p_desc = doc.add_paragraph()
                        p_desc.paragraph_format.space_before = Pt(0)
                        p_desc.paragraph_format.space_after = Pt(2)
                        p_desc.paragraph_format.keep_with_next = True
                        run_desc = p_desc.add_run(company_desc[0].strip())
                        run_desc.font.name = font_name
                        run_desc.font.size = Pt(9.5)
                        run_desc.italic = True
                        run_desc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    if meta:
                        p_meta = doc.add_paragraph()
                        p_meta.paragraph_format.space_before = Pt(0)
                        p_meta.paragraph_format.space_after = Pt(3)
                        p_meta.paragraph_format.keep_with_next = True
                        run_meta = p_meta.add_run(meta[0].strip())
                        run_meta.font.name = font_name
                        run_meta.font.size = Pt(9.5)
                        run_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                        run_meta.italic = True

                    for li in bullets:
                        p_bullet = doc.add_paragraph(style='List Bullet')
                        p_bullet.paragraph_format.space_before = Pt(0)
                        p_bullet.paragraph_format.space_after = Pt(2)
                        p_bullet.paragraph_format.left_indent = Inches(0.2)
                        
                        li_children = li.xpath('./* | text()')
                        for node in li_children:
                            if isinstance(node, str):
                                text = node
                                if text:
                                    run = p_bullet.add_run(text)
                                    run.font.name = font_name
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                            else:
                                text = node.text_content()
                                if text:
                                    run = p_bullet.add_run(text)
                                    run.font.name = font_name
                                    run.font.size = Pt(10)
                                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                    if node.tag == 'strong':
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

                elif el_class == 'skills-grid':
                    skills_lines = el.xpath('.//div[@class="skills-line"]')
                    for line in skills_lines:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        
                        strong_el = line.find('strong')
                        if strong_el is not None:
                            run_bold = p.add_run(strong_el.text_content().strip() + " ")
                            run_bold.bold = True
                            run_bold.font.name = font_name
                            run_bold.font.size = Pt(10)
                            run_bold.font.color.rgb = primary_color_rgb
                            
                            rest_text = line.text_content()[len(strong_el.text_content()):].strip()
                            run_norm = p.add_run(rest_text)
                            run_norm.font.name = font_name
                            run_norm.font.size = Pt(10)
                            run_norm.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                        else:
                            run_norm = p.add_run(line.text_content().strip())
                            run_norm.font.name = font_name
                            run_norm.font.size = Pt(10)
                            run_norm.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

                elif el_class == 'projects-grid':
                    cards = el.xpath('.//div[@class="project-card"]')
                    if cards:
                        table = doc.add_table(rows=(len(cards) + 1) // 2, cols=2)
                        table.autofit = False
                        
                        widths = [Inches(3.75), Inches(3.75)]
                        for row in table.rows:
                            for idx, width in enumerate(widths):
                                row.cells[idx].width = width
                        
                        for idx, card in enumerate(cards):
                            row_idx = idx // 2
                            col_idx = idx % 2
                            cell = table.cell(row_idx, col_idx)
                            
                            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                            
                            card_inner = cell.add_paragraph()
                            card_inner.paragraph_format.space_before = Pt(6)
                            card_inner.paragraph_format.space_after = Pt(6)
                            
                            title_el = card.xpath('.//div[@class="project-title"]/text()')
                            desc_el = card.xpath('.//div[@class="project-desc"]/text()')
                            stack_el = card.xpath('.//div[@class="project-stack"]')
                            
                            if title_el:
                                run_title = card_inner.add_run(title_el[0].strip() + "\\n")
                                run_title.bold = True
                                run_title.font.name = font_name
                                run_title.font.size = Pt(10.5)
                                run_title.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                                
                            if desc_el:
                                run_desc = card_inner.add_run(desc_el[0].strip() + "\\n")
                                run_desc.font.name = font_name
                                run_desc.font.size = Pt(9.5)
                                run_desc.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                
                            if stack_el:
                                stack_text = stack_el[0].text_content().strip()
                                run_stack = card_inner.add_run(stack_text)
                                run_stack.font.name = font_name
                                run_stack.font.size = Pt(8.5)
                                run_stack.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
                                run_stack.italic = True

                elif el_class == 'education-bar':
                    left_text = el.xpath('.//span[@class="left"]//text()')
                    right_text = el.xpath('.//span[@class="right"]/text()')
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    
                    pPr = p._p.get_or_add_pPr()
                    if is_modern:
                        shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F0F6F9"/>')
                        pPr.append(shd)
                    
                    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                     rf'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{border_hex}"/>'
                                     r'</w:pBdr>')
                    pPr.append(pBdr)
                    
                    left_str = "".join(left_text).strip()
                    right_str = right_text[0].strip() if right_text else ""
                    
                    combined_str = f"{left_str}   |   {right_str}"
                    run = p.add_run(combined_str)
                    run.font.name = font_name
                    run.font.size = Pt(10)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

                elif el_class == 'flex-row-container':
                    table = doc.add_table(rows=1, cols=2)
                    table.autofit = False
                    
                    widths = [Inches(3.75), Inches(3.75)]
                    for row in table.rows:
                        for idx, width in enumerate(widths):
                            row.cells[idx].width = width
                    
                    cols = el.xpath('.//div[@class="flex-row-col"]')
                    for col_idx, col in enumerate(cols):
                        cell = table.cell(0, col_idx)
                        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                        
                        sub_title = col.xpath('.//div[@class="section-title"]')
                        if sub_title:
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(8)
                            p.paragraph_format.space_after = Pt(4)
                            run = p.add_run(sub_title[0].text_content().strip().upper())
                            run.font.name = font_name
                            run.font.size = Pt(11)
                            run.bold = True
                            run.font.color.rgb = primary_color_rgb
                            add_p_border_bottom(p, color_hex=border_hex, size="6")
                            
                        cert_list = col.xpath('.//div[@class="cert-list"]')
                        if cert_list:
                            lines = [l.strip() for l in cert_list[0].text_content().split('\\n') if l.strip()]
                            for line in lines:
                                p = cell.add_paragraph()
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(2)
                                run = p.add_run(line)
                                run.font.name = font_name
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                                
                        ach_list = col.xpath('.//ul[@class="achievements-list"] | .//ul')
                        if ach_list:
                            for li in ach_list[0].xpath('.//li'):
                                p = cell.add_paragraph(style='List Bullet')
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.space_after = Pt(2)
                                p.paragraph_format.left_indent = Inches(0.2)
                                run = p.add_run(li.text_content().strip())
                                run.font.name = font_name
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                
                i += 1

        # Save document
        output_path = os.path.join(DOWNLOADS_DIR, f"{filename_prefix}.docx")
        doc.save(output_path)
        print(f"Generated Word Docx: {output_path}")

def generate_pdfs():
    """Launches Playwright headless Chromium to open page and save PDF files."""
    print("Launching Playwright for PDF generation...")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        
        abs_html_path = os.path.abspath(HTML_FILE)
        page.goto(f"file:///{abs_html_path}")
        
        page.wait_for_selector(".sidebar-title")
        
        for ver_id, filename_prefix in VERSIONS.items():
            version_tab_id, style_id = ver_id.split("-")
            
            page.click(f"#btn-{version_tab_id}")
            page.click(f"#btn-style-{style_id}")
            page.wait_for_timeout(500)
            
            output_path = os.path.join(DOWNLOADS_DIR, f"{filename_prefix}.pdf")
            page.pdf(
                path=output_path,
                format="Letter",
                print_background=True,
                margin={"top": "0.45in", "bottom": "0.45in", "left": "0.45in", "right": "0.45in"}
            )
            print(f"Generated PDF: {output_path}")
            
        browser.close()

if __name__ == "__main__":
    print("=== Starting Resume Generation ===")
    setup_directories()
    
    print("\n--- 1. Extracting LaTeX ---")
    extract_latex()
    
    print("\n--- 2. Generating Word Documents ---")
    generate_word_docs()
    
    print("\n--- 3. Generating PDFs via Playwright ---")
    generate_pdfs()
    
    print("\n=== Generation Complete! ===")
